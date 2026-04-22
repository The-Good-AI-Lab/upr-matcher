from collections.abc import Iterable
from pathlib import Path
from typing import TYPE_CHECKING, Any, Final

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastembed import TextEmbedding
from openai import OpenAI
from pypdf import PdfReader

try:
    from pydantic_ai import Agent
    from pydantic_ai.models.openai import OpenAIChatModel
    from pydantic_ai.output import PromptedOutput
    from pydantic_ai.providers.openrouter import OpenRouterProvider
except ImportError:  # pragma: no cover - handled at runtime
    Agent = None  # type: ignore[assignment]
    OpenAIModel = None  # type: ignore[assignment]
    OpenRouterProvider = None  # type: ignore[assignment]

if TYPE_CHECKING:
    from pydantic_ai.models.openai import OpenAIChatModel as OpenAIChatModelType
else:  # pragma: no cover
    OpenAIChatModelType = Any

from .settings import Settings

settings = Settings()


def _normalize_cell_text(value: str) -> str:
    return value.strip().replace("\xa0", " ")


def _require_openrouter_key() -> str:
    if not settings.openrouter_api_key:
        raise ValueError(
            "OPENROUTER_API_KEY is not set. Please export it in your environment before running this command."
        )
    return settings.openrouter_api_key


def get_openrouter_client() -> OpenAI:
    return OpenAI(
        base_url=settings.agent_base_url,
        api_key=_require_openrouter_key(),
    )


_text_embedder: TextEmbedding | None = None
_chat_model: OpenAIChatModelType | None = None


def get_text_embedder() -> TextEmbedding:
    global _text_embedder
    if _text_embedder is None:
        _text_embedder = TextEmbedding(model_name=settings.embedding_model)
    return _text_embedder


def prompt_openrouter(prompt: str) -> str:
    client = get_openrouter_client()
    response = client.chat.completions.create(model=settings.model, messages=[{"role": "user", "content": prompt}])
    return response.choices[0].message.content


def message_openrouter(messages: list[dict]) -> str:
    client = get_openrouter_client()
    response = client.chat.completions.create(model=settings.model, messages=messages)
    return response.choices[0].message.content


def chat_with_openrouter(system_prompt: str, user_prompt: str) -> str:
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": [{"type": "text", "text": user_prompt}]},
    ]
    return message_openrouter(messages)


def _get_chat_model() -> OpenAIChatModelType:
    if OpenAIChatModel is None or Agent is None or OpenRouterProvider is None:
        raise ImportError(
            "pydantic-ai is required for structured outputs. Install it or run `uv sync` to fetch dependencies."
        )
    global _chat_model
    if _chat_model is None:
        provider = OpenRouterProvider(api_key=_require_openrouter_key())
        _chat_model = OpenAIChatModel(settings.model, provider=provider)
    return _chat_model


def structured_chat_openrouter(system_prompt: str, user_prompt: str, response_format: type[Any]) -> Any:
    agent = Agent(
        model=_get_chat_model(),
        system_prompt=system_prompt,
    )
    result = agent.run_sync(user_prompt, output_type=PromptedOutput(response_format))
    return result.output


def _table_to_rows(table: Table) -> Iterable[str]:
    # Assuming the first row contains headers
    headers = [_normalize_cell_text(cell.text) for cell in table.rows[0].cells]
    for row in table.rows[1:]:
        cells = [_normalize_cell_text(cell.text) for cell in row.cells]
        if len(cells) != len(headers):
            continue
        yield "\n\n".join(f"{headers[i]}\n{cells[i]}" for i in range(len(headers)))


def _table_to_text(table: Table) -> str:
    return "\n\n".join(_table_to_rows(table))


def _table_to_json(table: Table) -> list[dict]:
    # Assuming the first row contains headers
    headers = [_normalize_cell_text(cell.text) for cell in table.rows[0].cells]
    json_data: list[dict[str, str]] = []
    current_theme: dict[str, str] = {}
    for row in table.rows[1:]:
        cells = [_normalize_cell_text(cell.text) for cell in row.cells]
        # if the row has 1 cell then it's probably a title row, skip it keep as metadata
        # Example: "Theme: Legal & institutional reform"
        if cells[0].lower().startswith("theme"):
            cell_text = cells[0]
            if ":" in cell_text:
                key, value = cell_text.split(":", 1)
                current_theme = {key.strip(): value.strip()}
            continue
        if len(cells) != len(headers):
            continue
        row_data = {headers[i]: cells[i] for i in range(len(headers))}
        if not any(value for value in row_data.values()):
            continue
        if current_theme:
            row_data = {**current_theme, **row_data}
        json_data.append(row_data)
    return json_data


def docx_tables_to_json(path: Path | str) -> list[dict[str, str]]:
    document = Document(Path(path))
    rows: list[dict[str, str]] = []
    for table in document.tables:
        rows.extend(_table_to_json(table))
    return rows


def _docx_blocks_in_order(doc: Document) -> Iterable[str]:
    # check if the doc has only tables
    if all(el.tag.rsplit("}", 1)[-1] == "tbl" for el in doc.element.body.iterchildren()):
        for table in doc.tables:
            table_text = _table_to_text(table)
            if table_text.strip():
                yield table_text
        return
    # otherwise yield paragraphs and tables in order
    for el in doc.element.body.iterchildren():
        tag = el.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph_text = Paragraph(el, doc).text
            if paragraph_text.strip():
                yield paragraph_text
        elif tag == "tbl":
            table_text = _table_to_text(Table(el, doc))
            if table_text.strip():
                yield table_text


def read_text_file(path: Path | str) -> str:
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        reader: Final = PdfReader(p.open("rb"))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if ext == ".docx":
        doc: Final = Document(p)
        return "\n\n".join(block.strip() for block in _docx_blocks_in_order(doc) if block.strip())
    if ext in {".txt", ".md"}:
        return p.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported file type: {ext}")
